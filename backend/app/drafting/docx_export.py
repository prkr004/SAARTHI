"""Programmatic DOCX assembly for strict drafting payloads.

The drafting pipeline now enforces extraction-first generation. The LLM returns
schema-validated fields and this module applies a deterministic document layout.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from backend.app.drafting.schema import (
    AdvisoryDraft,
    CircularDraft,
    DraftDocument,
    PressReleaseDraft,
    canonicalize_document_type,
    validate_draft_payload,
)

try:  # pragma: no cover - import guard for environments that have not installed the dependency yet
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - handled at runtime
    Document = None  # type: ignore[assignment]
    WD_TABLE_ALIGNMENT = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    parse_xml = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    Inches = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]
    _DOCX_IMPORT_ERROR = exc
else:
    _DOCX_IMPORT_ERROR = None

_TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"


def template_path(document_type: str) -> Path:
    """Return the on-disk template path for a supported document type."""

    normalized = canonicalize_document_type(document_type)
    return _TEMPLATE_DIR / f"{normalized}.txt"


def load_template_text(document_type: str) -> str:
    """Load the plain-text template associated with a document type."""

    path = template_path(document_type)
    if not path.exists():
        raise FileNotFoundError(f"Template file not found for document type '{document_type}'.")
    return path.read_text(encoding="utf-8")


def _coerce_document_draft(document_data: DraftDocument | dict[str, Any]) -> DraftDocument:
    if isinstance(document_data, (CircularDraft, AdvisoryDraft, PressReleaseDraft)):
        return document_data
    return validate_draft_payload(document_data)


def _require_docx() -> None:
    if _DOCX_IMPORT_ERROR is not None:
        raise RuntimeError(
            "python-docx is required for DOCX export. Install the project dependencies first."
        ) from _DOCX_IMPORT_ERROR


def _set_paragraph_spacing(paragraph, *, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    format_ = paragraph.paragraph_format
    format_.space_before = Pt(before)
    format_.space_after = Pt(after)
    format_.line_spacing = line


def _set_font(run, *, size: int = 11, bold: bool = False, italic: bool = False, name: str = "Times New Roman") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)  # type: ignore[union-attr]
    except Exception:
        pass


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 11, alignment=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    _set_font(run, size=size, bold=bold)
    paragraph.alignment = alignment if alignment is not None else getattr(WD_ALIGN_PARAGRAPH, "LEFT", None)
    _set_paragraph_spacing(paragraph, after=0)


def _hide_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    borders = parse_xml(
        r"<w:tblBorders xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
        r"<w:top w:val='nil'/><w:left w:val='nil'/><w:bottom w:val='nil'/><w:right w:val='nil'/>"
        r"<w:insideH w:val='nil'/><w:insideV w:val='nil'/></w:tblBorders>"
    )
    tbl_pr.append(borders)


def _configure_document(document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)


def _add_letterhead(document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("BANKING REGULATORY AUTHORITY")
    _set_font(run, size=14, bold=True)
    _set_paragraph_spacing(paragraph, after=12)


def _add_header_table(
    document,
    *,
    left_top: str,
    right_top: str,
    left_bottom: str,
    right_bottom: str = "",
) -> None:
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _hide_table_borders(table)

    _set_cell_text(table.cell(0, 0), left_top, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(table.cell(0, 1), right_top, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(table.cell(1, 0), left_bottom, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(table.cell(1, 1), right_bottom, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    spacer = document.add_paragraph()
    _set_paragraph_spacing(spacer, after=4)


def _add_subject_line(document, subject: str) -> None:
    paragraph = document.add_paragraph()
    prefix = paragraph.add_run("Subject: ")
    _set_font(prefix, bold=True)
    content = paragraph.add_run(subject)
    _set_font(content, bold=True)
    _set_paragraph_spacing(paragraph, after=10)


def _add_section_heading(document, heading: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(heading)
    _set_font(run, bold=True)
    _set_paragraph_spacing(paragraph, before=4, after=4)


def _add_justified_paragraph(document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    _set_font(run)
    _set_paragraph_spacing(paragraph)


def _add_bulleted_items(document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        _set_font(run)
        _set_paragraph_spacing(paragraph, after=4)


def _add_numbered_items(document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        _set_font(run)
        _set_paragraph_spacing(paragraph, after=4)


def _add_right_aligned_signoff(document, authority: str) -> None:
    spacer = document.add_paragraph()
    _set_paragraph_spacing(spacer, after=8)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(authority)
    _set_font(run, bold=True)
    _set_paragraph_spacing(paragraph, after=0)


def _buffer_from_document(document) -> BytesIO:
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def generate_circular_docx(data: CircularDraft) -> BytesIO:
    """Assemble a strict circular layout using typed extraction data."""

    _require_docx()
    document = Document()
    _configure_document(document)
    _add_letterhead(document)
    _add_header_table(
        document,
        left_top=f"Reference No: {data.reference_number}",
        right_top=f"Date: {data.date}",
        left_bottom=f"To: {data.addressee}",
    )
    _add_subject_line(document, data.subject)

    _add_section_heading(document, "Highlights")
    _add_bulleted_items(document, data.highlights)

    _add_section_heading(document, "Background / Reference")
    _add_justified_paragraph(document, data.background_context)

    _add_section_heading(document, "Operational Guidelines")
    _add_numbered_items(document, data.operational_directives)

    _add_section_heading(document, "Compliance / Action")
    _add_justified_paragraph(document, data.compliance_warning)

    _add_right_aligned_signoff(document, data.issuing_authority)
    return _buffer_from_document(document)


def generate_advisory_docx(data: AdvisoryDraft) -> BytesIO:
    """Assemble a strict advisory layout using typed extraction data."""

    _require_docx()
    document = Document()
    _configure_document(document)
    _add_letterhead(document)
    _add_header_table(
        document,
        left_top=f"Priority: {data.priority_level}",
        right_top=f"Date: {data.date}",
        left_bottom=f"Target Audience: {data.target_audience}",
    )
    _add_subject_line(document, data.subject)

    _add_section_heading(document, "Issue Description")
    _add_justified_paragraph(document, data.issue_description)

    _add_section_heading(document, "Mitigating Actions")
    _add_numbered_items(document, data.mitigating_actions)

    _add_section_heading(document, "Reporting Mechanism")
    _add_justified_paragraph(document, data.reporting_mechanism)

    _add_right_aligned_signoff(document, data.issuing_authority)
    return _buffer_from_document(document)


def generate_press_release_docx(data: PressReleaseDraft) -> BytesIO:
    """Assemble a strict press release layout using typed extraction data."""

    _require_docx()
    document = Document()
    _configure_document(document)
    _add_letterhead(document)

    immediate = document.add_paragraph()
    immediate.alignment = WD_ALIGN_PARAGRAPH.CENTER
    immediate_run = immediate.add_run("FOR IMMEDIATE RELEASE")
    _set_font(immediate_run, bold=True)
    _set_paragraph_spacing(immediate, after=8)

    _add_header_table(
        document,
        left_top=f"Date: {data.date}",
        right_top=f"Dateline: {data.dateline}",
        left_bottom="",
    )

    headline = document.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline_run = headline.add_run(data.headline)
    _set_font(headline_run, size=13, bold=True)
    _set_paragraph_spacing(headline, after=10)

    _add_section_heading(document, "Lead")
    _add_justified_paragraph(document, data.lead_paragraph)

    _add_section_heading(document, "Body")
    for paragraph_text in data.body_paragraphs:
        _add_justified_paragraph(document, paragraph_text)

    _add_section_heading(document, "About")
    _add_justified_paragraph(document, data.boilerplate_about)

    _add_section_heading(document, "Media Contact")
    media = document.add_paragraph()
    media_run = media.add_run(data.media_contact)
    _set_font(media_run)
    _set_paragraph_spacing(media)

    return _buffer_from_document(document)


def create_docx(document_data: DraftDocument | dict[str, Any], *, bank_name: str = "SAARTHI") -> bytes:
    """Dispatch strict typed payload to the correct deterministic DOCX assembler."""

    _ = bank_name
    _require_docx()
    draft = _coerce_document_draft(document_data)

    if isinstance(draft, CircularDraft):
        return generate_circular_docx(draft).getvalue()
    if isinstance(draft, AdvisoryDraft):
        return generate_advisory_docx(draft).getvalue()
    return generate_press_release_docx(draft).getvalue()


def export_docx_preview(document_data: DraftDocument | dict[str, Any]) -> str:
    """Render a text preview from strict typed payload fields."""

    draft = _coerce_document_draft(document_data)

    if isinstance(draft, CircularDraft):
        lines = [
            f"[Circular] {draft.subject}",
            f"Ref: {draft.reference_number} | Date: {draft.date}",
            f"To: {draft.addressee}",
            "Highlights:",
            *[f"- {item}" for item in draft.highlights],
            "Operational Directives:",
            *[f"- {item}" for item in draft.operational_directives],
        ]
        return "\n".join(lines)

    if isinstance(draft, AdvisoryDraft):
        lines = [
            f"[Advisory] {draft.subject}",
            f"Priority: {draft.priority_level} | Date: {draft.date}",
            f"Target Audience: {draft.target_audience}",
            "Mitigating Actions:",
            *[f"- {item}" for item in draft.mitigating_actions],
        ]
        return "\n".join(lines)

    lines = [
        f"[Press Release] {draft.headline}",
        f"Date: {draft.date} | Dateline: {draft.dateline}",
        "Body:",
        *[f"- {item}" for item in draft.body_paragraphs],
    ]
    return "\n".join(lines)


def export_docx(document_data: DraftDocument | dict[str, Any]) -> bytes:
    """Backward-compatible wrapper for the DOCX exporter."""

    return create_docx(document_data)
